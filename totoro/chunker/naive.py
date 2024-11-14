#!/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File    :   pdf.py
@Time    :   2024/09/06 16:37:31
@Desc    :
'''

import json
import re
from io import BytesIO
from pydoc import doc
from timeit import default_timer as timer
from typing import List

import chardet
from docx import Document
from tika import parser

from totoro.parser import (DocxParser, ExcelParser, HtmlParser,
                           PdfParser, PlainParser)
from totoro.pb import doc_pb2
from totoro.nlp.tokenizer import DocTokenizer
from totoro.utils.encoder import ModelEncoder
from totoro.utils.logger import logger
from totoro.utils.utils import get_file_type, get_snowflake_id

from .chunker import ChunkBuilder

# from totoro.utils import num_tokens_from_string


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        callback(msg="Naive PDF OCR is running...")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        print(f"self.boxes after __images__ is {len(self.boxes)} ")
        callback(msg="OCR finished")
        logger.info("OCR({}~{}): {}".format(
            from_page, to_page, timer() - start))

        start = timer()
        logger.debug("start to layout analysis")
        self._layouts_rec(zoomin)
        print(f"self.boxes after _layouts_rec is {len(self.boxes)} ")
        logger.debug("layout analysis finished")
        callback(0.63, "Layout analysis finished.")
        self._table_transformer_job(zoomin)
        print(f"self.boxes after _table_transformer_job is {len(self.boxes)} ")
        callback(0.65, "Table analysis finished.")
        self._text_merge()
        print(f"self.boxes after _text_merge is {len(self.boxes)} ")
        callback(0.67, "Text merging finished")
        tbls = self._extract_table_figure(True, zoomin, True, True)
        print(f"self.boxes after _extract_table_figure is {len(self.boxes)} ")
        # self._naive_vertical_merge()
        print(f"self.boxes after _naive_vertical_merge is {len(self.boxes)} ")
        self._concat_downward()
        print(f"self.boxes after _concat_downward is {len(self.boxes)} ")
        # self._filter_forpages()
        print(f"self.boxes after _filter_forpages is {len(self.boxes)} ")

        logger.info("layouts: {}".format(timer() - start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return [(line, "") for line in lines if line], tbls


class NaiveChunkBuilder(ChunkBuilder):
    def __init__(self):
        super().__init__()
        print("call NaiveChunkBuilder.__init__")
        self.tokenizer = DocTokenizer()
        self.encoder = ModelEncoder()

    def naive_merge(self, sections, chunk_token_num=128, delimiter="\n。；！？"):
        if not sections:
            return []
        if isinstance(sections[0], type("")):
            sections = [(s, "") for s in sections]
        cks = [""]
        tk_nums = [0]

        def add_chunk(t, pos):
            nonlocal cks, tk_nums, delimiter
            tnum = self.encoder.num_tokens_from_string(t)
            if tnum < 8:
                pos = ""
            if tk_nums[-1] > chunk_token_num:
                if t.find(pos) < 0:
                    t += pos
                cks.append(t)
                tk_nums.append(tnum)
            else:
                if cks[-1].find(pos) < 0:
                    t += pos
                cks[-1] += t
                tk_nums[-1] += tnum

        for sec, pos in sections:
            add_chunk(sec, pos)
            continue

        return cks

    def chunk(self, filename, doc_name, binary=None, from_page=0, to_page=100000,
              lang="Chinese", callback=None, **kwargs) -> List[doc_pb2.Doc]:
        """
            Supported file formats are docx, pdf, excel, txt.
            This method apply the naive ways to chunk files.
            Successive text will be sliced into pieces using 'delimiter'.
            Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
        """
        eng = lang.lower() == "english"  # is_english(cks)
        parser_config = kwargs.get(
            "parser_config", {
                "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": True})
        # doc = {
        #     "doc_name_keyword": filename,
        #     "title_tokens": self.tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
        # }
        doc = doc_pb2.Doc(
            doc_name_keyword=doc_name,
            title_tokens=self.tokenizer.tokenize(
                re.sub(
                    r"\.[a-zA-Z]+$",
                    "",
                    doc_name)))

        doc.title_small_tokens = self.tokenizer.fine_grained_tokenize(
            doc.title_tokens)
        res = []
        pdf_parser = None
        sections = []
        file = binary or filename
        file_type = get_file_type(file)

        if file_type == "docx":
            callback(0.1, "Start to parse.")
            sections, tbls = Docx()(filename, binary)
            res = self.tokenizer.tokenize_table(tbls, doc, eng)
            callback(1.00, "Finish parsing.")
        elif file_type == "pdf":
            pdf_parser = Pdf(
            ) if getattr(parser_config, "layout_recognize", True) else PlainParser()
            sections, tbls = pdf_parser(filename if not binary else binary,
                                        from_page=from_page, to_page=to_page, callback=callback)

            res = self.tokenizer.tokenize_table(tbls, doc, eng)

        elif file_type == "xlsx":
            callback(0.1, "Start to parse.")
            excel_parser = ExcelParser()
            sections = [(line, "")
                        for line in excel_parser.html(binary) if line]
        elif re.search(r"\.(txt|md|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt)$", filename, re.IGNORECASE):
            callback(0.1, "Start to parse.")
            txt = ""
            if binary:
                detected = chardet.detect(binary)

                encoding = detected.get("encoding", "utf-8")
                txt = binary.decode(encoding, errors="ignore")
            else:
                with open(filename, "r") as f:
                    while True:
                        line = f.readline()
                        if not line:
                            break
                        txt += line
            sections = []
            for sec in txt.split("\n"):
                if self.encoder.num_tokens_from_string(sec) > 10 * parser_config.get("chunk_token_num", 128):
                    sections.append((sec[:int(len(sec) / 2)], ""))
                    sections.append((sec[int(len(sec) / 2):], ""))
                else:
                    sections.append((sec, ""))
            callback(1.00, "Finish parsing.")
        elif file_type == "html":
            callback(0.1, "Start to parse.")
            sections = HtmlParser()(filename, binary)
            sections = [(line, "") for line in sections if line]
            callback(1.00, "Finish parsing.")
        elif file_type == "doc" or re.search(r"\.doc$", filename, re.IGNORECASE):
            callback(0.1, "Start to parse.")
            binary = BytesIO(binary)
            doc_parsed = parser.from_buffer(binary)
            sections = doc_parsed['content'].split('\n')
            sections = [(line, "") for line in sections if line]
            callback(1.00, "Finish parsing.")
        else:
            raise NotImplementedError(
                f"file type {file_type} not supported yet(pdf, xlsx, doc, docx, txt supported)")
        st = timer()
        chunks = self.naive_merge(
            sections, getattr(parser_config, "chunk_token_num", 128), getattr(parser_config, "delimiter", "\n!?。；！？"))
        res.extend(self.tokenizer.tokenize_chunks(
            chunks, doc, eng, pdf_parser))
        logger.info("naive_merge({}): {}".format(filename, timer() - st))
        return res
